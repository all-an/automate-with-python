import docx

address = 'C:\\Users\\Public\\code\\automate-with-python\\013-excel-word-pdf'

file1 = '\\demo.docx'

dir1 = address + file1

document = docx.Document(dir1)

print(type(document))

print(document.paragraphs[0].text)

p2 = document.paragraphs[1]

print(p2.runs[0].text)

print(p2.runs[1].bold) # True

print(p2.runs[1].underline)

p2.runs[1].underline = True

document.add_paragraph('hello new paragraph')

document.save(address + '\\demo2.docx')